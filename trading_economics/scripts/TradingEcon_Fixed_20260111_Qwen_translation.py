# -*- coding: utf-8 -*-
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.edge.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.edge.options import Options
from bs4 import BeautifulSoup
import pyperclip
import time
import tkinter as tk
from tkinter import messagebox
from tkinter import simpledialog
from docx import Document
from datetime import datetime
import os
import tempfile
import shutil

# 添加requests库
import requests
from urllib.parse import urljoin

# 导入Dashscope相关模块
try:
    import dashscope
    print("DEBUG: dashscope successfully imported")
    from dashscope import Generation  # 只导入Generation，不需要Message

    dashscope_available = True
    print("DEBUG: dashscope library is available")
except ImportError as e:
    print(f"DEBUG: ImportError when importing dashscope: {e}")
    print("警告: dashscope 库未安装，翻译功能将被跳过。如需使用翻译功能，请运行: pip install dashscope")
    dashscope_available = False
except Exception as e:
    print(f"DEBUG: Unexpected error when importing dashscope: {e}")
    print("警告: dashscope 库出现问题，翻译功能将被跳过。")
    dashscope_available = False


def safe_click_element(driver, locator, timeout=30):
    """安全点击元素，带重试机制和多种策略"""
    try:
        wait = WebDriverWait(driver, timeout)
        # 等待元素可点击
        element = wait.until(EC.element_to_be_clickable(locator))
        
        # 尝试滚动到元素可见区域
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)
        time.sleep(1)  # 等待滚动完成
        
        # 尝试普通点击
        element.click()
        return True
    except Exception as e:
        print(f"普通点击元素失败: {e}")
        
        try:
            # 尝试使用 ActionChains 点击
            element = wait.until(EC.presence_of_element_located(locator))
            actions = ActionChains(driver)
            actions.move_to_element(element).click().perform()
            return True
        except Exception as e2:
            print(f"ActionChains 点击也失败: {e2}")
            
            try:
                # 最后的备选方案：使用 JavaScript 点击
                element = wait.until(EC.presence_of_element_located(locator))
                # 再次尝试滚动到视图
                driver.execute_script("arguments[0].scrollIntoView(true);", element)
                time.sleep(1)
                driver.execute_script("arguments[0].click();", element)
                return True
            except Exception as e3:
                print(f"JavaScript 点击也失败: {e3}")
                
                # 额外尝试：隐藏可能遮挡的元素
                try:
                    driver.execute_script("""
                        var elements = document.querySelectorAll('nav, header, .navbar, .fixed-top');
                        for(var i = 0; i < elements.length; i++) {
                            elements[i].style.display = 'none';
                        }
                    """)
                    time.sleep(1)
                    element = wait.until(EC.element_to_be_clickable(locator))
                    element.click()
                    return True
                except:
                    print("所有点击方法都失败了")
                    return False


def retry_on_timeout(func, max_retries=3, delay=5):
    """重试装饰器，处理超时错误"""
    for attempt in range(max_retries):
        try:
            return func()
        except Exception as e:
            if "timeout" in str(e).lower() and attempt < max_retries - 1:
                print(f"超时错误，第 {attempt + 1} 次重试，等待 {delay} 秒...")
                time.sleep(delay)
                continue
            else:
                raise e


def safe_calendar_click(driver, wait, time_range):
    """安全的日历点击函数，优化版本"""
    try:
        # 等待下拉按钮可用并滚动到视图
        dropdown_button = wait.until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, "button.btn.btn-outline-secondary.btn-calendar"))
        )
        # 滚动到元素可见
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", dropdown_button)
        time.sleep(1)
        
        # 尝试多种点击方法
        try:
            dropdown_button.click()
        except:
            # 如果普通点击失败，尝试JavaScript点击
            driver.execute_script("arguments[0].click();", dropdown_button)
        
        # 等待选项出现并点击
        xpath_str = f"//li[contains(@class, 'dropdown-item te-c-option')]//a[@onclick=\"setCalendarRange('{time_range}')\"]"
        this_week_option = wait.until(
            EC.element_to_be_clickable((By.XPATH, xpath_str))
        )
        # 再次确保元素可见
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", this_week_option)
        time.sleep(1)
        
        try:
            this_week_option.click()
        except:
            # 如果普通点击失败，尝试JavaScript点击
            driver.execute_script("arguments[0].click();", this_week_option)
    except Exception as e:
        print(f"日历点击失败: {e}")
        # 使用 JavaScript 作为备用方案
        driver.execute_script(f"setCalendarRange('{time_range}');")


def extract_description_from_html(html_content):
    """从HTML内容中提取描述文本"""
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 查找具有特定属性的div元素
    specified_element = soup.select_one('div[role="tabpanel"] h2[id="description"]')
    if specified_element:
        return specified_element.get_text(strip=True)
    
    # 如果没找到，尝试其他可能的选择器
    fallback_selectors = [
        'h2#description',
        'div[role="tabpanel"] h2',
        'h2',
        '.content p',
        'p'
    ]
    
    for selector in fallback_selectors:
        element = soup.select_one(selector)
        if element:
            return element.get_text(strip=True)
    
    return ""


def get_content_from_url(url, session=None):
    """使用requests获取URL内容"""
    if session is None:
        session = requests.Session()
    
    # 设置请求头，模拟真实浏览器
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36 Edg/91.0.864.59',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        'Accept-Language': 'en-US,en;q=0.5',
        'Accept-Encoding': 'gzip, deflate',
        'Connection': 'keep-alive',
        'Upgrade-Insecure-Requests': '1',
    }
    
    try:
        full_url = urljoin("https://tradingeconomics.com", url)
        response = session.get(full_url, headers=headers, timeout=15)
        response.raise_for_status()
        
        # 提取描述内容
        content = extract_description_from_html(response.text)
        return content if content else "未找到描述内容"
    except requests.exceptions.RequestException as e:
        print(f"获取URL失败 {url}: {e}")
        return f"获取失败: {e}"


# 创建临时用户目录
user_data_dir = tempfile.mkdtemp()

# 设置Edge选项以提高性能
options = Options()
options.add_argument(f"--user-data-dir={user_data_dir}")  # 指定唯一用户目录
options.add_argument("--disable-cleaning-data-cache")  # 禁用自动清理数据缓存
options.add_argument("--disable-images")  # 禁用图片加载以提高速度
options.add_argument("--disable-javascript")  # 如果不需要JS可禁用（谨慎使用）
options.add_argument("--disable-plugins")  # 禁用插件
options.add_argument("--no-sandbox")  # 避免一些沙盒问题
options.add_argument("--disable-dev-shm-usage")  # 解决资源限制问题
options.add_argument("--disable-extensions")  # 禁用扩展
options.add_argument("--disable-web-security")  # 禁用网络安全检查
options.add_argument("--allow-running-insecure-content")  # 允许不安全内容运行
options.add_experimental_option('excludeSwitches', ['enable-logging'])  # 禁用日志
options.add_experimental_option('useAutomationExtension', False)  # 禁用自动化扩展
options.add_argument("--disable-background-timer-throttling")  # 禁用后台定时器节流
options.add_argument("--disable-renderer-backgrounding")  # 禁用渲染器后台处理
options.add_argument("--disable-backgrounding-occluded-windows")  # 禁用隐藏窗口的后台处理
options.add_argument("--window-size=1920,1080")  # 设置窗口大小，确保元素可见

# 使用Service对象初始化Edge Web驱动
service = Service('C:\\msedgedriver\\msedgedriver.exe')
driver = webdriver.Edge(service=service, options=options)  # 传入options参数

# 设置各种超时
driver.set_page_load_timeout(180)  # 页面加载超时
driver.implicitly_wait(30)  # 隐式等待

# 打开目标网址
print("正在打开TradingEconomics日历页面...")
driver.get('https://tradingeconomics.com/calendar')

# 增加显式等待时间
wait = WebDriverWait(driver, 30)

# 等待页面基本元素加载
WebDriverWait(driver, 30).until(
    EC.presence_of_element_located((By.CSS_SELECTOR, "button.btn.btn-outline-secondary.btn-calendar"))
)
print("页面已加载完成")

# 创建Tkinter窗口以获取用户输入
root = tk.Tk()
root.withdraw()  # 隐藏主窗口
time_range = tk.simpledialog.askstring("请输入所需新闻的时间",
                                       "-1为昨天，-2为上周，1为今天，2为明天，3为本周，4为下周，5为本月，6为下周")
root.destroy()

# 使用改进后的日历点击函数
safe_calendar_click(driver, wait, time_range)

# 创建Tkinter窗口以获取用户输入
root = tk.Tk()
root.withdraw()  # 隐藏主窗口
news_importance = tk.simpledialog.askstring("选择需要获取的重要程度",
                                            "请输入新闻的重要程度，1为普通，2为重要，3为十分重要）:")
root.destroy()

# 找到并点击用于打开下拉菜单的按钮
importance_dropdown_locator = (By.ID, 'ctl00_ContentPlaceHolder1_ctl02_Button1')
safe_click_element(driver, importance_dropdown_locator)

# 在下拉菜单中找到并点击指定重要性的选项
importance_option_xpath = f"//a[@onclick=\"setCalendarImportance('{news_importance}');\"]"
importance_option_locator = (By.XPATH, importance_option_xpath)
safe_click_element(driver, importance_option_locator)

# 点击 'Countries' 下拉菜单
countries_dropdown_locator = (By.CSS_SELECTOR, "button.btn.btn-outline-secondary.btn-calendar[onclick='toggleMainCountrySelection();']")
safe_click_element(driver, countries_dropdown_locator)

# 等待下拉菜单展开
time.sleep(1)

# 点击 'Clear' 按钮
clear_button_locator = (By.CSS_SELECTOR, "a.btn.btn-outline-secondary.te-c-option[onclick='clearSelection();']")
safe_click_element(driver, clear_button_locator)

# 定义要选中的国家列表
countries_to_select = "li.te-c-option.te-c-option-chn, li.te-c-option.te-c-option-emu, li.te-c-option.te-c-option-usa"

# 直接选择多个国家
country_elements = wait.until(EC.presence_of_all_elements_located((By.CSS_SELECTOR, countries_to_select)))

# 使用JavaScript点击每个国家元素，避免元素被遮挡的问题
for country_element in country_elements:
    driver.execute_script("arguments[0].click();", country_element)

# 最后点击Save
save_button_locator = (By.CSS_SELECTOR, "a.btn.btn-success.te-c-option")
safe_click_element(driver, save_button_locator)

# 获取页面源代码
page_source = driver.page_source

# 使用BeautifulSoup解析HTML
soup = BeautifulSoup(page_source, 'html.parser')

# 获取当前日期，格式为"YYYY-MM-DD"
current_date_str = datetime.now().strftime("%Y-%m-%d")

# 使用列表推导式来筛选出不包含特定字符串的URL
urls = []

# 遍历每个<thead>元素
for thead in soup.find_all('thead', class_='hidden-head'):
    # 从<thead>中提取日期字符串
    date_str = thead.find('th', colspan="3").get_text(strip=True)
    # 将日期字符串转换为datetime对象
    event_date = datetime.strptime(date_str, '%A %B %d %Y')

    # 将当前日期的字符串转换为datetime对象
    current_date = datetime.strptime(current_date_str, '%Y-%m-%d')

    # 检查事件日期是否晚于当前日期
    if event_date > current_date:
        continue  # 如果是，则跳过该<thead>下的<tbody>内容

    # 处理<thead>下的<tbody>中的<tr>元素
    tbody = thead.find_next_sibling('tbody')
    if tbody:
        for tr in tbody.find_all('tr', attrs={'data-url': True, 'data-event': True}):
            tr_soup = BeautifulSoup(str(tr), 'html.parser')
            actual_span_element = tr_soup.select_one('span#actual')

            # 检查span id="actual">的内容
            actual_span_text = actual_span_element.text.strip() if actual_span_element else ""

            # 只有当span id="actual">不是"$ 0"时，才添加URL
            if actual_span_text != "$ 0":
                span_element = tr_soup.select_one('span.calendar-date-2')

                # 获取span标签的文本内容，如果标签不存在则设置为空字符串
                span_text = span_element.text if span_element else ""

                if ("speech" not in tr['data-event'].lower()) and \
                        ("calendar" not in tr['data-url'].lower()) and \
                        ("ecb monetary policy meeting accounts" not in tr['data-event'].lower()) and \
                        ("your_condition_here" not in span_text.lower()):  # 根据您的实际需求更新your_condition_here"
                    urls.append(tr['data-url'])

# 输出所有找到的URL
full_url = None
for url in urls:
    print(url)

total_urls = len(urls)
opened_urls = 0

print(f"Total URLs to traverse: {total_urls}")

element_contents = []

# 使用requests会话来复用连接，提高效率
session = requests.Session()

for url in urls:
    try:
        full_url = "https://tradingeconomics.com" + url
        print(f"Trying to get: {full_url}")
        
        # 使用requests获取内容，而不是Selenium
        content = get_content_from_url(url, session)
        if content and content != "未找到描述内容":
            element_contents.append(f"\"{content}\"")
        else:
            print(f"无法找到描述内容: {full_url}")
            continue

    except Exception as e:
        print(f"An error occurred while getting content from {full_url}: {e}")
        continue

# 关闭Selenium浏览器，因为我们已经获取了需要的URL列表
driver.quit()

for content in element_contents:
    print(content)

print(f"Total URLs traversed: {total_urls}")
print(f"Total URLs processed: {len(element_contents)}")

joined_contents = "\n\n".join(element_contents)
pyperclip.copy(joined_contents)

# 创建一个新的Document对象
doc = Document()

# 添加原文内容
doc.add_paragraph("英文原文:")
doc.add_paragraph(joined_contents)

# 检查并创建目标目录 - 改为保存到 C:\\Onedrive\\TradingEconTransRawData
target_directory = "C:\\Onedrive\\TradingEconTransRawData"
directory = os.path.dirname(target_directory)
if not os.path.exists(directory):
    os.makedirs(directory, exist_ok=True)

# 生成不冲突的文件名
def get_available_filename(base_path):
    if not os.path.exists(base_path):
        return base_path
    
    name, ext = os.path.splitext(base_path)
    counter = 1
    while True:
        new_path = f"{name}_{counter}{ext}"
        if not os.path.exists(new_path):
            return new_path
        counter += 1

file_path = f"C:\\Onedrive\\TradingEconTransRawData\\{current_date_str}pyTED.doc"
file_path = get_available_filename(file_path)

# 如果dashscope可用，使用Qwen API进行国别分析总结
if dashscope_available:
    # 从环境变量获取API密钥
    api_key = os.environ.get('DASHSCOPE_API_KEY_Qwen')
    if api_key:
        dashscope.api_key = api_key
        print("DEBUG: API key found, attempting country analysis...")
        
        try:
            # 使用Qwen进行国别分析总结
            analysis_response = Generation.call(
                model="qwen-max",  # 更强大的模型，能够更好地进行分析
                messages=[
                    {'role': 'system',
                     'content': '专业经济分析师。'},
                    {'role': 'user',
                     'content': f'请基于以下数据，对上周中、美、欧三个国家和地区的经济数据进行深度分析，分析结构应包括：\n1. 关键指标表现：GDP、工业生产、零售销售、就业、通胀等核心数据\n2. 最新数据变化：突出最新公布数据的趋势和变化特点\n3. 趋势走向：分析当前经济运行的主要特征，如内外需表现差异等\n4. 重要观察点：深入分析数据背后的驱动因素和结构性问题\n5. 风险与前瞻：识别潜在风险点并提供前瞻性判断\n\n要求：突出结构性问题、政策影响、风险点和增长前景。每个国家用一段连贯的文字表述，避免简单罗列数据。数据：{joined_contents}'}
                ]
            )
            
            print(f"DEBUG: Analysis API response status: {analysis_response.status_code}")
            
            if analysis_response.status_code == 200:
                # 获取分析总结内容
                if hasattr(analysis_response, "output") and \
                   hasattr(analysis_response.output, "choices") and \
                   analysis_response.output.choices:
                    analysis_content = analysis_response.output.choices[0].message.content
                    print("DEBUG: Analysis content retrieved successfully")
                elif hasattr(analysis_response, "output") and \
                     hasattr(analysis_response.output, "text") and \
                     analysis_response.output.text:
                    analysis_content = analysis_response.output.text
                    print("DEBUG: Analysis content retrieved from text field")
                else:
                    analysis_content = "国别分析总结未能生成"
                    print(f"DEBUG: No content found in analysis response: {analysis_response}")
                
                # 添加国别分析到文档
                doc.add_paragraph("\n\n国别经济数据分析总结:")
                doc.add_paragraph(analysis_content)
                print("\n\n国别经济数据分析总结:\n" + analysis_content)
                print("国别分析总结完成")
            else:
                print(f"国别分析失败: {analysis_response.code}, {analysis_response.message}")
                # 如果分析失败，仍然继续后续步骤
                # 添加错误信息到文档
                doc.add_paragraph("\n\n国别经济数据分析总结:")
                doc.add_paragraph(f"国别分析失败: {analysis_response.code}, {analysis_response.message}")
        except Exception as e:
            print(f"国别分析过程中出现错误: {e}")
            # 添加错误信息到文档
            doc.add_paragraph("\n\n国别经济数据分析总结:")
            doc.add_paragraph(f"国别分析过程中出现错误: {e}")
    else:
        print("未找到环境变量 DASHSCOPE_API_KEY_Qwen，跳过国别分析功能")
        # 即使没有API密钥，也在文档中记录这一情况
        doc.add_paragraph("\n\n国别经济数据分析总结:")
        doc.add_paragraph("未找到API密钥，跳过国别分析功能")
else:
    print("由于dashscope库未安装，跳过国别分析功能")
    # 即使库不可用，也在文档中记录这一情况
    doc.add_paragraph("\n\n国别经济数据分析总结:")
    doc.add_paragraph("由于dashscope库未安装，跳过国别分析功能")

# 如果dashscope可用，使用Qwen API进行翻译
if dashscope_available:
    # 从环境变量获取API密钥
    api_key = os.environ.get('DASHSCOPE_API_KEY_Qwen')
    if api_key:
        dashscope.api_key = api_key
        print("DEBUG: API key found, attempting translation...")

        try:
            # 使用Qwen进行翻译 - 修正消息格式
            response = Generation.call(
                model="qwen-flash",  # 适用于快速翻译的模型
                messages=[  # 使用字典格式的消息
                    {'role': 'system', 'content': '你是一个专业的财经翻译专家，擅长将经济数据和分析报告从英文翻译成中文。请保持原文的专业术语准确性、数据精确性，并维持正式的财经报告语体。'},
                    {'role': 'user', 'content': f'请将以下英文经济数据和分析报告翻译成中文，要求：\n1. 保持所有数字、百分比、日期等数据的精确性\n2. 使用标准的财经术语，如GDP、PMI、CPI等保持英文缩写，记得先写中文然后紧跟在括号里的英文缩写\n3. 保持正式的财经报告语体\n4. 注意术语翻译：flash→初值, advance estimate→初步估算, yoy→同比, mom→环比。若原文没有明确环比和同比，再翻译的时候请根据上下文，区分环比和同比\n5. 确保译文逻辑清晰，符合中文表达习惯\n\n经济数据内容：{joined_contents}'}
                ]
            )

            print(f"DEBUG: Translation API response status: {response.status_code}")
            
            if response.status_code == 200:
                # 检查response.output.choices是否存在且非空，否则检查response.output.text
                if hasattr(response, "output") and \
                   hasattr(response.output, "choices") and \
                   response.output.choices:
                    translated_content = response.output.choices[0].message.content
                    doc.add_paragraph("\n\n中文翻译:")
                    doc.add_paragraph(translated_content)
                    print("\n\n中文翻译:\n" + translated_content)
                    print("翻译成功完成")
                elif hasattr(response, "output") and \
                     hasattr(response.output, "text") and \
                     response.output.text:
                    # 如果choices为空，尝试使用text字段
                    translated_content = response.output.text
                    doc.add_paragraph("\n\n中文翻译:")
                    doc.add_paragraph(translated_content)
                    print("\n\n中文翻译:\n" + translated_content)
                    print("翻译成功完成")
                else:
                    print(f"DEBUG: No content found in translation response: {response}")
                    doc.add_paragraph("\n\n中文翻译:")
                    doc.add_paragraph("翻译服务返回内容格式异常")
            else:
                print(f"翻译失败: {response.code}, {response.message}")
                # 如果翻译失败，仍然保存原文档
                # 添加错误信息到文档
                doc.add_paragraph("\n\n中文翻译:")
                doc.add_paragraph(f"翻译失败: {response.code}, {response.message}")
        except Exception as e:
            print(f"翻译过程中出现错误: {e}")
            # 添加错误信息到文档
            doc.add_paragraph("\n\n中文翻译:")
            doc.add_paragraph(f"翻译过程中出现错误: {e}")
    else:
        print("未找到环境变量 DASHSCOPE_API_KEY_Qwen，请设置API密钥以使用翻译功能")
        # 即使没有API密钥，也在文档中记录这一情况
        doc.add_paragraph("\n\n中文翻译:")
        doc.add_paragraph("未找到API密钥，跳过翻译功能")
else:
    print("由于dashscope库未安装，跳过翻译功能")
    # 即使库不可用，也在文档中记录这一情况
    doc.add_paragraph("\n\n中文翻译:")
    doc.add_paragraph("由于dashscope库未安装，跳过翻译功能")

# 保存文档（无论是否进行了翻译）
doc.save(file_path)

print("程序执行完成")

# 创建提示窗口
root = tk.Tk()
root.withdraw()

# 弹出消息框
messagebox.showinfo("任务完成", "所有任务已经完成，即将打开相关文档")

# 销毁Tkinter窗口
root.destroy()

# 等待一段时间确保所有操作完成
time.sleep(2)

# 关闭浏览器驱动
driver.quit()

# 删除临时目录
try:
    shutil.rmtree(user_data_dir)
except Exception as e:
    print(f"删除临时目录时出错: {e}")

# 使用默认应用程序打开文件
os.startfile(file_path)