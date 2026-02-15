import sqlite3
import os
import json
import csv
import pandas as pd
import logging
from datetime import datetime
from typing import List, Dict, Optional

class DataStorage:
    def __init__(self, db_path: str = 'data/db/wechat.db'):
        self.db_path = os.path.abspath(db_path)
        self.logger = logging.getLogger(__name__)
        self._init_db()
    
    def _init_db(self):
        """初始化数据库"""
        try:
            # 确保数据库目录存在
            os.makedirs(os.path.dirname(self.db_path), exist_ok=True)
            
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # 创建accounts表
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS accounts (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    account_name TEXT UNIQUE,
                    account_id TEXT,
                    last_update TEXT,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            # 创建articles表
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS articles (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    account_id TEXT,
                    account_name TEXT,
                    title TEXT,
                    summary TEXT,
                    content TEXT,
                    url TEXT UNIQUE,
                    publish_time TEXT,
                    cover_image TEXT,
                    reading_count INTEGER DEFAULT 0,
                    like_count INTEGER DEFAULT 0,
                    crawl_time TEXT,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (account_id) REFERENCES accounts(account_id)
                )
            ''')
            
            # 检查是否需要添加summary字段（兼容旧数据库）
            try:
                cursor.execute('SELECT summary FROM articles LIMIT 1')
            except sqlite3.OperationalError:
                # summary字段不存在，添加它
                self.logger.info("正在升级数据库，添加summary字段...")
                cursor.execute('ALTER TABLE articles ADD COLUMN summary TEXT')
                self.logger.info("数据库升级完成")
            
            conn.commit()
            conn.close()
            self.logger.info(f"数据库初始化成功: {self.db_path}")
        except Exception as e:
            self.logger.error(f"数据库初始化失败: {str(e)}")
    
    def _connect(self):
        """连接数据库"""
        return sqlite3.connect(self.db_path)
    
    def save_account(self, account_info: Dict):
        """保存公众号信息"""
        try:
            conn = self._connect()
            cursor = conn.cursor()
            
            # 检查是否已存在
            cursor.execute('''
                SELECT id FROM accounts WHERE account_name = ?
            ''', (account_info.get('name'),))
            existing = cursor.fetchone()
            
            if existing:
                # 更新
                cursor.execute('''
                    UPDATE accounts SET account_id = ?, last_update = ?
                    WHERE account_name = ?
                ''', (account_info.get('id', ''), datetime.now().isoformat(), account_info.get('name')))
            else:
                # 插入
                cursor.execute('''
                    INSERT INTO accounts (account_name, account_id, last_update)
                    VALUES (?, ?, ?)
                ''', (account_info.get('name'), account_info.get('id', ''), datetime.now().isoformat()))
            
            conn.commit()
            conn.close()
            self.logger.info(f"保存公众号成功: {account_info.get('name')}")
        except Exception as e:
            self.logger.error(f"保存公众号失败: {str(e)}")
    
    def save_articles(self, articles: List[Dict]):
        """批量保存文章"""
        if not articles:
            return
        
        try:
            conn = self._connect()
            cursor = conn.cursor()
            
            # 批量插入
            batch_size = 100
            for i in range(0, len(articles), batch_size):
                batch = articles[i:i+batch_size]
                for article in batch:
                    try:
                        # 检查是否已存在
                        cursor.execute('''
                            SELECT id FROM articles WHERE url = ?
                        ''', (article.get('url', ''),))
                        existing = cursor.fetchone()
                        
                        if not existing:
                            cursor.execute('''
                                INSERT INTO articles (
                                    account_id, account_name, title, summary, content, url, 
                                    publish_time, cover_image, reading_count, like_count, crawl_time
                                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                            ''', (
                                article.get('account_id', ''),
                                article.get('account_name', ''),
                                article.get('title', ''),
                                article.get('summary', ''),
                                article.get('content', ''),
                                article.get('url', ''),
                                article.get('publish_time', ''),
                                article.get('cover_image', ''),
                                article.get('reading_count', 0),
                                article.get('like_count', 0),
                                article.get('crawl_time', datetime.now().isoformat())
                            ))
                    except Exception as e:
                        self.logger.warning(f"保存文章失败: {str(e)}")
                        continue
            
            conn.commit()
            conn.close()
            self.logger.info(f"保存文章成功: {len(articles)} 篇")
        except Exception as e:
            self.logger.error(f"批量保存文章失败: {str(e)}")
    
    def get_accounts(self) -> List[Dict]:
        """获取所有公众号"""
        try:
            conn = self._connect()
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('SELECT * FROM accounts ORDER BY created_at DESC')
            accounts = [dict(row) for row in cursor.fetchall()]
            
            conn.close()
            return accounts
        except Exception as e:
            self.logger.error(f"获取公众号失败: {str(e)}")
            return []
    
    def get_account_by_name(self, name: str) -> Optional[Dict]:
        """根据名称获取公众号"""
        try:
            conn = self._connect()
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('SELECT * FROM accounts WHERE account_name = ?', (name,))
            row = cursor.fetchone()
            
            conn.close()
            return dict(row) if row else None
        except Exception as e:
            self.logger.error(f"获取公众号失败: {str(e)}")
            return None
    
    def get_last_update(self, account_id: str) -> Optional[str]:
        """获取公众号上次更新时间"""
        try:
            conn = self._connect()
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT last_update FROM accounts WHERE account_id = ?
            ''', (account_id,))
            result = cursor.fetchone()
            
            conn.close()
            return result[0] if result else None
        except Exception as e:
            self.logger.error(f"获取上次更新时间失败: {str(e)}")
            return None
    
    def update_last_update(self, account_id: str):
        """更新公众号上次更新时间"""
        try:
            conn = self._connect()
            cursor = conn.cursor()
            
            cursor.execute('''
                UPDATE accounts SET last_update = ? WHERE account_id = ?
            ''', (datetime.now().isoformat(), account_id))
            
            conn.commit()
            conn.close()
        except Exception as e:
            self.logger.error(f"更新上次更新时间失败: {str(e)}")
    
    def get_articles(self, account_name: Optional[str] = None, limit: int = 100) -> List[Dict]:
        """获取文章"""
        try:
            conn = self._connect()
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            if account_name:
                cursor.execute('''
                    SELECT * FROM articles WHERE account_name = ?
                    ORDER BY publish_time DESC LIMIT ?
                ''', (account_name, limit))
            else:
                cursor.execute('''
                    SELECT * FROM articles
                    ORDER BY publish_time DESC LIMIT ?
                ''', (limit,))
            
            articles = [dict(row) for row in cursor.fetchall()]
            conn.close()
            return articles
        except Exception as e:
            self.logger.error(f"获取文章失败: {str(e)}")
            return []
    
    def get_article_by_url(self, url: str) -> Optional[Dict]:
        """根据URL获取文章"""
        try:
            conn = self._connect()
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('SELECT * FROM articles WHERE url = ?', (url,))
            row = cursor.fetchone()
            
            conn.close()
            return dict(row) if row else None
        except Exception as e:
            self.logger.error(f"获取文章失败: {str(e)}")
            return None
    
    def save_article(self, article: Dict):
        """保存单篇文章"""
        try:
            conn = self._connect()
            cursor = conn.cursor()
            
            # 检查是否已存在
            cursor.execute('''
                SELECT id FROM articles WHERE url = ?
            ''', (str(article.get('url', '')),))
            existing = cursor.fetchone()
            
            # 确保所有字段都是字符串类型
            account_id = str(article.get('account_id', ''))
            account_name = str(article.get('account_name', ''))
            title = str(article.get('title', ''))
            summary = str(article.get('summary', ''))
            content = str(article.get('content', ''))
            url = str(article.get('url', ''))
            publish_time = str(article.get('publish_time', ''))
            cover_image = str(article.get('cover_image', ''))
            reading_count = int(article.get('reading_count', 0))
            like_count = int(article.get('like_count', 0))
            crawl_time = str(article.get('crawl_time', datetime.now().isoformat()))
            
            if not existing:
                # 插入新文章
                cursor.execute('''
                    INSERT INTO articles (
                        account_id, account_name, title, summary, content, url, 
                        publish_time, cover_image, reading_count, like_count, crawl_time
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    account_id, account_name, title, summary, content, url, 
                    publish_time, cover_image, reading_count, like_count, crawl_time
                ))
                conn.commit()
                self.logger.info(f"保存文章成功: {title[:30]}...")
            else:
                # 检查是否需要更新content字段
                cursor.execute('''
                    SELECT content FROM articles WHERE url = ?
                ''', (url,))
                existing_content = cursor.fetchone()
                
                if existing_content and not existing_content[0] and content:
                    # 如果现有content为空，且新content不为空，则更新
                    cursor.execute('''
                        UPDATE articles 
                        SET content = ?, crawl_time = ? 
                        WHERE url = ?
                    ''', (content, crawl_time, url))
                    conn.commit()
                    self.logger.info(f"更新文章内容: {title[:30]}...")
                else:
                    self.logger.info(f"文章已存在，跳过: {article.get('title', '')[:30]}...")
            
            conn.close()
        except Exception as e:
            self.logger.error(f"保存文章失败: {str(e)}")
    
    def get_all_articles(self) -> List[Dict]:
        """获取所有文章"""
        try:
            conn = self._connect()
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT * FROM articles
                ORDER BY publish_time DESC
            ''')
            
            articles = [dict(row) for row in cursor.fetchall()]
            conn.close()
            return articles
        except Exception as e:
            self.logger.error(f"获取所有文章失败: {str(e)}")
            return []
    
    def get_latest_article_by_account(self, account_name: str) -> Optional[Dict]:
        """获取指定公众号的最新文章"""
        try:
            conn = self._connect()
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT * FROM articles 
                WHERE account_name = ?
                ORDER BY publish_time DESC
                LIMIT 1
            ''', (account_name,))
            
            row = cursor.fetchone()
            conn.close()
            return dict(row) if row else None
        except Exception as e:
            self.logger.error(f"获取最新文章失败: {str(e)}")
            return None
    
    def article_exists(self, url: str) -> bool:
        """检查文章是否已存在"""
        try:
            conn = self._connect()
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT id FROM articles WHERE url = ?
            ''', (url,))
            
            exists = cursor.fetchone() is not None
            conn.close()
            return exists
        except Exception as e:
            self.logger.error(f"检查文章是否存在失败: {str(e)}")
            return False
    
    def export_to_csv(self, account_name: Optional[str] = None, filename: Optional[str] = None, articles: Optional[List[Dict]] = None):
        """导出为CSV"""
        try:
            # 如果没有传入articles，则从数据库获取
            if articles is None:
                articles = self.get_articles(account_name)
            
            if not articles:
                self.logger.warning("没有文章可导出")
                return None
            
            # 创建导出目录
            export_dir = 'data/processed'
            os.makedirs(export_dir, exist_ok=True)
            
            # 生成文件名
            if not filename:
                if account_name:
                    filename = f"{account_name}_{datetime.now().strftime('%Y%m%d')}.csv"
                else:
                    filename = f"all_articles_{datetime.now().strftime('%Y%m%d')}.csv"
            
            # 确保文件名有正确的扩展名
            if not filename.endswith('.csv'):
                filename += '.csv'
            
            filepath = os.path.join(export_dir, filename)
            
            # 导出为CSV
            df = pd.DataFrame(articles)
            df.to_csv(filepath, index=False, encoding='utf-8-sig')
            
            self.logger.info(f"导出CSV成功: {filepath}")
            return filepath
        except Exception as e:
            self.logger.error(f"导出CSV失败: {str(e)}")
            return None
    
    def export_to_json(self, account_name: Optional[str] = None, filename: Optional[str] = None, articles: Optional[List[Dict]] = None):
        """导出为JSON"""
        try:
            # 如果没有传入articles，则从数据库获取
            if articles is None:
                articles = self.get_articles(account_name)
            
            if not articles:
                self.logger.warning("没有文章可导出")
                return None
            
            # 创建导出目录
            export_dir = 'data/processed'
            os.makedirs(export_dir, exist_ok=True)
            
            # 生成文件名
            if not filename:
                if account_name:
                    filename = f"{account_name}_{datetime.now().strftime('%Y%m%d')}.json"
                else:
                    filename = f"all_articles_{datetime.now().strftime('%Y%m%d')}.json"
            
            # 确保文件名有正确的扩展名
            if not filename.endswith('.json'):
                filename += '.json'
            
            filepath = os.path.join(export_dir, filename)
            
            # 导出为JSON
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(articles, f, ensure_ascii=False, indent=2)
            
            self.logger.info(f"导出JSON成功: {filepath}")
            return filepath
        except Exception as e:
            self.logger.error(f"导出JSON失败: {str(e)}")
            return None
    
    def export_to_excel(self, account_name: Optional[str] = None, filename: Optional[str] = None, articles: Optional[List[Dict]] = None):
        """导出为Excel"""
        try:
            # 如果没有传入articles，则从数据库获取
            if articles is None:
                articles = self.get_articles(account_name)
            
            if not articles:
                self.logger.warning("没有文章可导出")
                return None
            
            # 创建导出目录
            export_dir = 'data/processed'
            os.makedirs(export_dir, exist_ok=True)
            
            # 生成文件名
            if not filename:
                if account_name:
                    filename = f"{account_name}_{datetime.now().strftime('%Y%m%d')}.xlsx"
                else:
                    filename = f"all_articles_{datetime.now().strftime('%Y%m%d')}.xlsx"
            
            # 确保文件名有正确的扩展名
            if not filename.endswith('.xlsx'):
                filename += '.xlsx'
            
            filepath = os.path.join(export_dir, filename)
            
            # 导出为Excel
            df = pd.DataFrame(articles)
            df.to_excel(filepath, index=False)
            
            self.logger.info(f"导出Excel成功: {filepath}")
            return filepath
        except Exception as e:
            self.logger.error(f"导出Excel失败: {str(e)}")
            return None
    
    def get_statistics(self) -> Dict:
        """获取统计信息"""
        try:
            conn = self._connect()
            cursor = conn.cursor()
            
            # 公众号数量
            cursor.execute('SELECT COUNT(*) FROM accounts')
            account_count = cursor.fetchone()[0]
            
            # 文章数量
            cursor.execute('SELECT COUNT(*) FROM articles')
            article_count = cursor.fetchone()[0]
            
            # 最近更新
            cursor.execute('SELECT MAX(last_update) FROM accounts')
            last_update = cursor.fetchone()[0]
            
            conn.close()
            
            return {
                'account_count': account_count,
                'article_count': article_count,
                'last_update': last_update,
                'db_path': self.db_path
            }
        except Exception as e:
            self.logger.error(f"获取统计信息失败: {str(e)}")
            return {}
    
    def export_to_word(self, account_name: Optional[str] = None, filename: Optional[str] = None):
        """导出为Word"""
        try:
            articles = self.get_articles(account_name)
            if not articles:
                self.logger.warning("没有文章可导出")
                return None
            
            # 创建导出目录
            export_dir = 'data/processed'
            os.makedirs(export_dir, exist_ok=True)
            
            # 生成文件名
            if not filename:
                if account_name:
                    filename = f"{account_name}_{datetime.now().strftime('%Y%m%d')}.docx"
                else:
                    filename = f"all_articles_{datetime.now().strftime('%Y%m%d')}.docx"
            
            filepath = os.path.join(export_dir, filename)
            
            # 导入python-docx库
            from docx import Document
            from docx.shared import Inches
            
            # 创建Word文档
            doc = Document()
            
            # 添加标题
            if account_name:
                doc.add_heading(f'{account_name} 公众号文章', 0)
            else:
                doc.add_heading('微信公众号文章合集', 0)
            
            doc.add_paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'文章数量: {len(articles)}')
            doc.add_page_break()
            
            # 为每篇文章创建章节
            for i, article in enumerate(articles, 1):
                # 添加文章标题
                doc.add_heading(f'{i}. {article.get("title", "无标题")}', level=1)
                
                # 添加文章元信息
                meta_info = []
                if article.get('account_name'):
                    meta_info.append(f'公众号: {article.get("account_name")}')
                if article.get('publish_time'):
                    meta_info.append(f'发布时间: {article.get("publish_time")}')
                if article.get('author'):
                    meta_info.append(f'作者: {article.get("author")}')
                if article.get('reading_count'):
                    meta_info.append(f'阅读量: {article.get("reading_count")}')
                if article.get('like_count'):
                    meta_info.append(f'点赞量: {article.get("like_count")}')
                if article.get('url'):
                    meta_info.append(f'原文链接: {article.get("url")}')
                
                if meta_info:
                    doc.add_paragraph(' | '.join(meta_info))
                
                # 添加文章内容
                content = article.get('content', '无内容')
                # 按换行符分割内容
                paragraphs = content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
                
                # 添加分页符
                if i < len(articles):
                    doc.add_page_break()
            
            # 保存文档
            doc.save(filepath)
            
            self.logger.info(f"导出Word成功: {filepath}")
            return filepath
        except ImportError:
            self.logger.error("导出Word失败: 缺少python-docx库，请安装: pip install python-docx")
            return None
        except Exception as e:
            self.logger.error(f"导出Word失败: {str(e)}")
            return None
    
    def clear_articles(self, account_name: Optional[str] = None):
        """清空文章数据"""
        try:
            conn = self._connect()
            cursor = conn.cursor()
            
            if account_name:
                cursor.execute('DELETE FROM articles WHERE account_name = ?', (account_name,))
            else:
                cursor.execute('DELETE FROM articles')
            
            conn.commit()
            conn.close()
            self.logger.info(f"清空文章数据成功")
        except Exception as e:
            self.logger.error(f"清空文章数据失败: {str(e)}")

# 测试代码
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    
    storage = DataStorage()
    
    # 测试获取统计信息
    stats = storage.get_statistics()
    print(f"统计信息: {stats}")
    
    # 测试获取公众号
    accounts = storage.get_accounts()
    print(f"公众号数量: {len(accounts)}")
