#!/usr/bin/env python3
"""
Word文档导出模块
支持将文章导出为Word格式，按公众号分类存储
"""

import os
import re
from datetime import datetime
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger(__name__)


class WordExporter:
    """Word文档导出类"""
    
    def __init__(self, base_dir: str = 'data/processed'):
        self.base_dir = base_dir
        os.makedirs(base_dir, exist_ok=True)
    
    def _sanitize_filename(self, filename: str) -> str:
        """清理文件名，移除非法字符"""
        # 移除或替换Windows文件名中的非法字符
        illegal_chars = '<>:"/\\|?*'
        for char in illegal_chars:
            filename = filename.replace(char, '_')
        # 限制文件名长度
        if len(filename) > 100:
            filename = filename[:100]
        return filename.strip()
    
    def _create_account_folder(self, account_name: str) -> str:
        """为公众号创建文件夹"""
        # 清理公众号名称作为文件夹名
        folder_name = self._sanitize_filename(account_name)
        folder_path = os.path.join(self.base_dir, folder_name)
        os.makedirs(folder_path, exist_ok=True)
        return folder_path
    
    def export_article_to_word(self, article: Dict, account_folder: str) -> Optional[str]:
        """
        将单篇文章导出为Word文档
        
        Args:
            article: 文章字典
            account_folder: 公众号文件夹路径
            
        Returns:
            导出的文件路径或None
        """
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置文档标题
            title = article.get('title', '无标题')
            
            # 添加标题
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加公众号信息
            account_name = article.get('account_name', '未知公众号')
            doc.add_paragraph(f"公众号：{account_name}")
            
            # 添加发布时间
            publish_time = article.get('publish_time', '未知时间')
            doc.add_paragraph(f"发布时间：{publish_time}")
            
            # 添加原文链接
            url = article.get('url', '')
            if url:
                doc.add_paragraph(f"原文链接：{url}")
            
            # 添加分隔线
            doc.add_paragraph("=" * 50)
            
            # 添加摘要
            summary = article.get('summary', '')
            if summary:
                doc.add_heading('摘要', level=2)
                doc.add_paragraph(summary)
            
            # 添加内容
            content = article.get('content', '')
            if content:
                doc.add_heading('正文', level=2)
                # 处理内容中的换行
                paragraphs = content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
            
            # 添加爬取信息
            doc.add_paragraph("=" * 50)
            crawl_time = article.get('crawl_time', datetime.now().isoformat())
            doc.add_paragraph(f"爬取时间：{crawl_time}")
            
            # 生成文件名
            # 使用标题和日期生成文件名
            safe_title = self._sanitize_filename(title)
            timestamp = datetime.now().strftime('%Y%m%d')
            filename = f"{timestamp}_{safe_title}.docx"
            
            # 确保文件名唯一
            filepath = os.path.join(account_folder, filename)
            counter = 1
            while os.path.exists(filepath):
                filename = f"{timestamp}_{safe_title}_{counter}.docx"
                filepath = os.path.join(account_folder, filename)
                counter += 1
            
            # 保存文档
            doc.save(filepath)
            logger.info(f"Word文档导出成功: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"导出Word文档失败: {str(e)}")
            return None
    
    def export_articles_by_account(self, articles: List[Dict]) -> Dict[str, List[str]]:
        """
        按公众号分类导出文章为Word文档
        
        Args:
            articles: 文章列表
            
        Returns:
            字典，key为公众号名称，value为导出的文件路径列表
        """
        exported_files = {}
        
        for article in articles:
            account_name = article.get('account_name', '未知公众号')
            
            # 创建公众号文件夹
            account_folder = self._create_account_folder(account_name)
            
            # 导出文章
            filepath = self.export_article_to_word(article, account_folder)
            
            if filepath:
                if account_name not in exported_files:
                    exported_files[account_name] = []
                exported_files[account_name].append(filepath)
                print(f"  [✓] {article.get('title', '无标题')[:40]}... -> Word")
        
        return exported_files
    
    def export_all_to_word(self, articles: List[Dict]) -> Dict[str, List[str]]:
        """
        将所有文章导出为Word文档（按公众号分类）
        
        Args:
            articles: 文章列表
            
        Returns:
            导出结果字典
        """
        print(f"\n{'='*60}")
        print("导出文章为Word文档")
        print(f"{'='*60}")
        print(f"共 {len(articles)} 篇文章需要导出")
        print(f"导出位置: {self.base_dir}")
        print(f"{'='*60}\n")
        
        exported = self.export_articles_by_account(articles)
        
        # 打印导出结果
        print(f"\n{'='*60}")
        print("导出完成")
        print(f"{'='*60}")
        
        total_files = 0
        for account_name, files in exported.items():
            print(f"\n公众号: {account_name}")
            print(f"  导出 {len(files)} 篇文章")
            total_files += len(files)
        
        print(f"\n总计: {total_files} 个Word文件")
        print(f"{'='*60}\n")
        
        return exported


if __name__ == "__main__":
    # 测试代码
    logging.basicConfig(level=logging.INFO)
    
    exporter = WordExporter()
    
    # 测试数据
    test_articles = [
        {
            'title': '测试文章1',
            'account_name': '测试公众号',
            'publish_time': '2026-02-14 10:00:00',
            'url': 'https://mp.weixin.qq.com/s/test1',
            'summary': '这是测试文章1的摘要',
            'content': '这是测试文章1的内容\n第二行内容',
            'crawl_time': datetime.now().isoformat()
        },
        {
            'title': '测试文章2',
            'account_name': '测试公众号',
            'publish_time': '2026-02-14 11:00:00',
            'url': 'https://mp.weixin.qq.com/s/test2',
            'summary': '这是测试文章2的摘要',
            'content': '这是测试文章2的内容',
            'crawl_time': datetime.now().isoformat()
        }
    ]
    
    result = exporter.export_all_to_word(test_articles)
    print("导出结果:", result)
