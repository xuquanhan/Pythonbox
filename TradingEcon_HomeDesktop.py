# -*- coding: utf-8 -*-

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.edge.service import Service
from selenium.webdriver.edge.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
import time
import tkinter as tk
from tkinter import messagebox, simpledialog
from docx import Document
from datetime import datetime
import os
import logging
import requests

# 配置日志记录
logging.basicConfig(filename='tradingecon.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logging.info("脚本开始执行")

# 配置 Edge 浏览器选项
edge_options = Options()

# AdGuard 插件路径（包含版本号文件夹）
adguard_extension_path = r"C:\msedgedriver\5.1.94_0"

# 添加 AdGuard 插件
edge_options.add_argument(f"--load-extension={adguard_extension_path}")

# 初始化 Edge WebDriver
service = Service(r'C:\msedgedriver\msedgedriver.exe')
driver = webdriver.Edge(service=service, options=edge_options)
logging.info("初始化 Edge WebDriver")

# 打开目标网址
driver.get('https://tradingeconomics.com/calendar')
logging.info("打开目标网址: https://tradingeconomics.com/calendar")

# 增加显式等待时间
wait = WebDriverWait(driver, 10)  # 增加等待时间

# 等待页面加载完成
wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "button.btn.btn-outline-secondary.btn-calendar")))

# 创建Tkinter窗口以获取用户输入
root = tk.Tk()
root.withdraw()  # 隐藏主窗口
time_range = tk.simpledialog.askstring("请输入所需新闻的时间段",
                                       "-1为昨天，-2为上周，1为今天，2为明天，3为本周，4为下周，5为本月，6为下月")
root.destroy()
logging.info(f"用户选择的时间段: {time_range}")

# 根据用户选择的时间段决定是否获取超过今日的数据
include_future_data = '1' if time_range in ['2', '3', '4', '5', '6'] else '0'
logging.info(f"是否获取超过今天的數據: {include_future_data}")

# 找到并点击用于打开下拉菜单的按钮
dropdown_button = wait.until(
    EC.element_to_be_clickable((By.CSS_SELECTOR, "button.btn.btn-outline-secondary.btn-calendar")))
dropdown_button.click()
logging.info("点击下拉菜单按钮")

# 在下拉菜单中找到并点击"用户选择的时间段"
xpath_str = f"//li[contains(@class, 'dropdown-item te-c-option')]//a[@onclick=\"setCalendarRange('{time_range}')\"]"
this_week_option = wait.until(EC.element_to_be_clickable((By.XPATH, xpath_str)))
this_week_option.click()
logging.info("点击用户选择的时间段")

# 找到并点击用于打开下拉菜单的按钮
dropdown_button = wait.until(EC.element_to_be_clickable((By.ID, 'ctl00_ContentPlaceHolder1_ctl02_Button1')))
dropdown_button.click()
logging.info("点击下拉菜单按钮")

# 创建Tkinter窗口以获取用户输入
root = tk.Tk()
root.withdraw()  # 隐藏主窗口
news_importance = tk.simpledialog.askstring("选择需要获取的重要程度",
                                            "请输入新闻的重要程度（1为普通，2为重要，3为十分重要）:")
root.destroy()
logging.info(f"用户选择的重要程度: {news_importance}")

# 在下拉菜单中找到并点击指定重要性的选项
importance_option_xpath = f"//a[@onclick=\"setCalendarImportance('{news_importance}');\"]"
importance_option = wait.until(
    EC.element_to_be_clickable((By.XPATH, importance_option_xpath)))
importance_option.click()
logging.info("点击指定重要性的选项")

# 点击 'Countries' 下拉菜单
countries_dropdown = wait.until(EC.element_to_be_clickable(
    (By.CSS_SELECTOR, "button.btn.btn-outline-secondary.btn-calendar[onclick='toggleMainCountrySelection();']")))
countries_dropdown.click()
logging.info("点击 Countries 下拉菜单")

# 隐藏广告元素
driver.execute_script("""
    const adElements = document.querySelectorAll('div#ad-slot.calendar-banner, div.modal-backdrop, div.loading-overlay');
    adElements.forEach(el => el.style.display = 'none');
""")
logging.info("隐藏广告元素")

# 截图验证广告是否隐藏
driver.save_screenshot('after_ad_hidden.png')
logging.info("保存广告隐藏后的截图")

# 点击 'Clear' 按钮
clear_button = wait.until(EC.element_to_be_clickable(
    (By.CSS_SELECTOR, "a.btn.btn-outline-secondary.te-c-option[onclick='clearSelection();']")))
clear_button.click()
logging.info("点击 Clear 按钮")

# 定义要选中的国家列表
countries_to_select = "li.te-c-option.te-c-option-chn, li.te-c-option.te-c-option-emu, li.te-c-option.te-c-option-usa"

# 直接选择多个国家
country_elements = wait.until(EC.presence_of_all_elements_located((By.CSS_SELECTOR, countries_to_select)))
for country_element in country_elements:
    country_element.click()
    logging.info(f"Clicked on country element: {country_element.text}")

# 最后点击Save
save_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "a.btn.btn-success.te-c-option")))
save_button.click()
logging.info("点击 Save 按钮")

# 等待页面跳转并加载完成
wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "tbody tr[data-url]")))
logging.info("等待页面跳转并加载完成")

# 获取页面源代码
page_source = driver.page_source
logging.debug(page_source)  # 使用 debug 级别记录页面源码

# 使用BeautifulSoup解析HTML
soup = BeautifulSoup(page_source, 'html.parser')

# 获取当前日期，格式为"YYYY-MM-DD"
current_date_str = datetime.now().strftime("%Y-%m-%d")

# 使用列表推导式来筛选出不包含特定字符串的URL
urls = []

# 遍历每个<thead>元素
for thead in soup.find_all('thead', class_='hidden-head'):
    # 从<thead>中提取日期字符串
    date_str = thead.find('th', colspan="3").get_text(strip=True)
    # 将日期字符串转换为datetime对象
    event_date = datetime.strptime(date_str, '%A %B %d %Y')

    # 将当前日期的字符串转换为datetime对象
    current_date = datetime.strptime(current_date_str, '%Y-%m-%d')

    # 根据用户选择调整日期过滤逻辑
    if include_future_data == '0' and event_date > current_date:
        logging.info(f"跳过日期晚于当前日期的事件: {date_str}")
        continue  # 如果是，则跳过该<thead>下的<tbody>内容

    # 处理该<thead>下的<tbody>中的<tr>元素
    tbody = thead.find_next_sibling('tbody')
    if tbody:
        for tr in tbody.find_all('tr', attrs={'data-url': True, 'data-event': True}):
            tr_soup = BeautifulSoup(str(tr), 'html.parser')
            actual_span_element = tr_soup.select_one('span#actual')

            # 检查span id="actual">的内容
            actual_span_text = actual_span_element.text.strip() if actual_span_element else ""
            logging.debug(f"实际值: {actual_span_text}")

            # 只有当span id="actual">不是"$ 0"时，才添加URL
            if actual_span_text != "$ 0":
                span_element = tr_soup.select_one('span.calendar-date-2')

                # 获取span标签的文本内容，如果标签不存在则设置为空字符串
                span_text = span_element.text if span_element else ""
                logging.debug(f"Span 文本: {span_text}")

                # 保持现有的过滤条件不变
                if ("speech" not in tr['data-event'].lower()) and \
                   ("ecb monetary policy meeting accounts" not in tr['data-event'].lower()):
                    urls.append(tr['data-url'])
                    logging.info(f"添加 URL: {tr['data-url']}")
                    print(f"添加 URL: {tr['data-url']}")  # 打印每条 URL
                else:
                    logging.info(f"跳过事件: {tr['data-event']}")
            else:
                logging.info(f"跳过实际值为 '$ 0' 的事件: {tr['data-event']}")

# 输出所有找到的URL
full_url = None
total_urls = len(urls)
opened_urls = 0

print(f"Total URLs to traverse: {total_urls}")

element_contents = []

for url in urls:
    try:
        # 跳过包含 /calendar 的URL
        if '/calendar' in url:
            logging.info(f"跳过包含 '/calendar' 的URL: {url}")
            continue

        full_url = "https://tradingeconomics.com" + url
        print(f"Trying to open: {full_url}")  # 打印每条 URL

        driver.get(full_url)
        opened_urls += 1

        time.sleep(0.1)

        specified_element = wait.until(
            EC.presence_of_element_located((By.CSS_SELECTOR, 'div[role="tabpanel"] h2[id="description"]')))
        content = specified_element.text
        element_contents.append((content, full_url))  # 将内容与URL保存为元组
        logging.info(f"提取到内容: {content[:100]}...")  # 记录部分内容
    except Exception as e:
        print(f"An error occurred while opening {full_url}: {e}")

# 打印总共获取的URL 数量
print(f"总共获取了{len(element_contents)} 条有效URL")

# 检查模型是否存在
try:
    response = requests.get("http://localhost:11434/api/tags")
    response.raise_for_status()
    models = [model['name'] for model in response.json().get('models', [])]
    if "deepseek-r1:7b" not in models:
        logging.error("模型 deepseek-r1:7b 未找到")
except Exception as e:
    logging.error(f"无法连接到Ollama服务: {e}")

# 修改翻译函数以记录详细信息
def translate_text(text, target_language='zh'):
    try:
        # 使用 Ollama 模型进行翻译，确保使用POST 方法
        response = requests.post("http://localhost:11434/api/generate", json={
            "model": "deepseek-r1:7b",
            "prompt": f"Translate the following text to {target_language}: {text}",
            "max_tokens": 2048
        })
        response.raise_for_status()
        result = response.json()["response"]
        logging.info(f"翻译结果: {result[:200]}...")  # 记录部分翻译结果
        return result
    except Exception as e:
        logging.error(f"翻译失败: {e}")
        return text

# 创建新的Document对象后添加文件路径定义
file_path = f"C:\\tradingecon\\{current_date_str}_pyTED.doc"  # 添加下划线分隔日期

# 确保目标目录存在
if not os.path.exists(os.path.dirname(file_path)):
    os.makedirs(os.path.dirname(file_path))

# 初始化 Document 对象
doc = Document()

# 遍历提取的内容并添加到文档中
for content, url in element_contents:
    doc.add_paragraph(f"URL: {url}")
    doc.add_paragraph(content)  # 直接添加原文，移除翻译功能
    doc.add_paragraph("\n")  # 添加段落分隔

# 保存文档到指定路径
doc.save(file_path)
logging.info(f"文档已保存至: {file_path}")

# 检查文件是否保存成功
if os.path.exists(file_path):
    logging.info("文件保存成功")
else:
    logging.error("文件保存失败")

# 创建提示窗口
root = tk.Tk()
root.withdraw()

# 弹出消息框
messagebox.showinfo("任务完成", "所有任务已经完成，即将打开相关文档。")

# 销毁Tkinter窗口
root.destroy()

# 关闭浏览器驱动
driver.quit()

# 使用默认应用程序打开文件
if os.path.exists(file_path):
    os.startfile(file_path)
else:
    logging.error("文件不存在，无法打开")